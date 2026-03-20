#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word文件读取工具
支持: .docx, .doc (Word 97-2003), WPS创建的文档
双击运行，选择文件后读取内容
"""

import os
import sys
import subprocess
import tempfile
import zipfile
import re
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from datetime import datetime
import time
import threading

# 尝试导入 python-docx
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，将无法读取 .docx 文件")

class WordReaderApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Word文档读取工具")
        self.root.geometry("1100x850")
        
        # 知识库路径 - 动态获取项目路径
        base_dir = os.path.dirname(os.path.abspath(__file__))
        self.library_path = os.path.join(base_dir, "errorLibrary", "Word_Library.md")
        
        # 当前文件路径
        self.current_file = None
        self.current_content = None
        
        # 精简Word相关
        self.clean_file = None
        self.clean_running = False
        
        # 校准相关
        self.calibration_running = False
        
        # 设置中文字体
        try:
            self.root.font = ("Microsoft YaHei", 10)
        except:
            pass
        
        self.create_widgets()
        self.root.mainloop()
    
    def create_widgets(self):
        """创建界面组件"""
        
        # ===== 顶部按钮区域 =====
        btn_frame = tk.Frame(self.root, pady=10)
        btn_frame.pack(fill=tk.X)
        
        self.select_btn = tk.Button(
            btn_frame, 
            text="📂 选择要被校准的文件", 
            command=self.select_file,
            font=("Microsoft YaHei", 12),
            padx=20,
            pady=5
        )
        self.select_btn.pack(side=tk.LEFT, padx=10)
        
        self.clear_btn = tk.Button(
            btn_frame,
            text="🗑️ 清空内容",
            command=self.clear_content,
            font=("Microsoft YaHei", 10),
            padx=15,
            pady=5
        )
        self.clear_btn.pack(side=tk.LEFT, padx=10)
        
        # ===== 主区域：左侧内容 + 右侧按钮 =====
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 左侧：内容显示区域
        left_frame = tk.Frame(main_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 文件信息标签
        self.file_label = tk.Label(
            left_frame, 
            text="请选择一个Word文件...",
            font=("Microsoft YaHei", 10),
            fg="gray"
        )
        self.file_label.pack(pady=5)
        
        # 内容显示区域
        self.text_area = scrolledtext.ScrolledText(
            left_frame,
            wrap=tk.WORD,
            font=("Microsoft YaHei", 11),
            padx=10,
            pady=10
        )
        self.text_area.pack(fill=tk.BOTH, expand=True)
        
        # 右侧：按钮区域
        right_frame = tk.Frame(main_frame, width=150)
        right_frame.pack(side=tk.RIGHT, fill=tk.Y, padx=(10, 0))
        right_frame.pack_propagate(False)
        
        # 右侧按钮 - 执行校准
        self.calibrate_btn = tk.Button(
            right_frame,
            text="🔧 执行校准",
            command=self.run_calibrate,
            font=("Microsoft YaHei", 11),
            padx=15,
            pady=10,
            bg="#4CAF50",
            fg="white",
            width=12
        )
        self.calibrate_btn.pack(pady=10)
        
        # 右侧按钮 - 知识库
        self.library_btn = tk.Button(
            right_frame,
            text="📚 知识库",
            command=self.open_library,
            font=("Microsoft YaHei", 11),
            padx=15,
            pady=10,
            width=12
        )
        self.library_btn.pack(pady=10)
        
        # 分隔线
        tk.Frame(right_frame, height=2, bg="gray").pack(fill=tk.X, padx=5, pady=10)
        
        # 右侧按钮 - 选择被精简的Word
        self.clean_select_btn = tk.Button(
            right_frame,
            text="📂 选择被精简的Word",
            command=self.select_clean_file,
            font=("Microsoft YaHei", 10),
            padx=10,
            pady=8,
            width=15
        )
        self.clean_select_btn.pack(pady=5)
        
        # 右侧按钮 - 执行精简
        self.clean_btn = tk.Button(
            right_frame,
            text="✂️ 执行精简",
            command=self.run_clean,
            font=("Microsoft YaHei", 11),
            padx=15,
            pady=8,
            bg="#9C27B0",
            fg="white",
            width=12
        )
        self.clean_btn.pack(pady=5)
        
        # 被精简文件路径
        self.clean_file_label = tk.Label(
            right_frame,
            text="未选择文件",
            font=("Microsoft YaHei", 8),
            fg="gray",
            wraplength=120
        )
        self.clean_file_label.pack(pady=5)
        
        # 设置按钮已隐藏
        # self.settings_btn = tk.Button(...)
        
        # ===== 底部：日志输出区域（可调整高度）=====
        # 使用 PanedWindow 实现可调整高度
        self.log_paned = tk.PanedWindow(self.root, orient=tk.VERTICAL, sashrelief=tk.RAISED, sashwidth=6)
        self.log_paned.pack(fill=tk.X, padx=10, pady=(5, 10))
        
        # 日志标题栏
        log_title_frame = tk.Frame(self.log_paned)
        self.log_paned.add(log_title_frame, height=30)
        
        log_title = tk.Label(
            log_title_frame,
            text="📋 日志输出 ▼ 拖动调整高度",
            font=("Microsoft YaHei", 9),
            anchor=tk.W
        )
        log_title.pack(side=tk.LEFT, padx=5)
        
        # 清空日志按钮
        tk.Button(
            log_title_frame,
            text="🗑️ 清空日志",
            command=self.clear_log,
            font=("Microsoft YaHei", 8),
            padx=10
        ).pack(side=tk.RIGHT, padx=5)
        
        # 日志文本区域（高度为原来的1.5倍，约12行）
        self.log_text = scrolledtext.ScrolledText(
            self.log_paned,
            wrap=tk.WORD,
            font=("Consolas", 9),
            height=12,
            bg="#1e1e1e",
            fg="#00ff00"
        )
        self.log_paned.add(self.log_text)
        
        # 状态栏
        self.status_label = tk.Label(
            self.root,
            text="就绪",
            font=("Microsoft YaHei", 9),
            fg="gray",
            anchor=tk.W
        )
        self.status_label.pack(fill=tk.X, padx=10, pady=(0, 5))
        
        # 绑定快捷键
        self.root.bind('<Control-o>', lambda e: self.select_file())
        self.root.bind('<Control-q>', lambda e: self.root.quit())
        
        # 初始化日志
        self.log("程序已启动")
        self.log(f"知识库路径: {self.library_path}")
    
    def log(self, message, tag=None):
        """输出日志"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_message = f"[{timestamp}] {message}\n"
        
        self.log_text.insert(tk.END, log_message)
        
        if tag:
            start_index = self.log_text.index(tk.END + "-2l")
            end_index = self.log_text.index(tk.END + "-1c")
            self.log_text.tag_config(tag, foreground=self._get_tag_color(tag))
            self.log_text.tag_add(tag, start_index, end_index)
        
        self.log_text.see(tk.END)
        self.root.update()
    
    def _get_tag_color(self, tag):
        """获取标签颜色"""
        colors = {
            "info": "#00ff00",
            "success": "#00FF00",
            "warning": "#FFA500",
            "error": "#FF0000",
            "highlight": "#FFFF00",
            "title": "#00FFFF",
            "result": "#FF69B4"
        }
        return colors.get(tag, "#00ff00")
    
    def log_replace(self, line_num, wrong, right):
        """输出替换日志（带颜色和行号）"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_message = f"[{timestamp}] 替换: 第{line_num}行 「{wrong}」 → 「{right}」\n"
        
        self.log_text.insert(tk.END, log_message)
        
        # 设置替换内容为黄色
        start_index = self.log_text.index(tk.END + "-2l")
        end_index = self.log_text.index(tk.END + "-1c")
        self.log_text.tag_config("replace", foreground="#FFFF00")
        self.log_text.tag_add("replace", start_index, end_index)
        
        self.log_text.see(tk.END)
        self.root.update()
    
    def clear_log(self):
        """清空日志"""
        self.log_text.delete(1.0, tk.END)
        self.log("日志已清空")
    
    def select_file(self):
        """选择文件"""
        file_types = [
            ("Word文档", "*.docx *.doc"),
            ("Word 2007+", "*.docx"),
            ("Word 97-2003", "*.doc"),
            ("所有文件", "*.*")
        ]
        
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=file_types,
            initialdir=os.path.expanduser("~/Desktop")
        )
        
        if filename:
            self.read_file(filename)
    
    def read_file(self, filepath):
        """读取文件内容"""
        self.current_file = filepath
        self.file_label.config(text=os.path.basename(filepath), fg="black")
        self.status_label.config(text="正在读取...")
        self.log(f"开始读取文件: {filepath}")
        self.text_area.delete(1.0, tk.END)
        
        self.root.update()
        
        try:
            ext = os.path.splitext(filepath)[1].lower()
            self.log(f"文件类型: {ext}")
            
            if ext == '.docx':
                content = self.read_docx(filepath)
                self.current_content = content
            elif ext == '.doc':
                content = self.read_doc(filepath)
                self.current_content = content
            else:
                content = f"不支持的文件格式: {ext}"
            
            self.text_area.insert(1.0, content)
            char_count = len(content)
            self.status_label.config(text=f"读取完成 - {char_count} 字符")
            self.log(f"读取完成: {char_count} 字符")
            
        except Exception as e:
            error_msg = f"读取失败: {str(e)}"
            self.text_area.insert(1.0, error_msg)
            self.status_label.config(text="读取失败", fg="red")
            self.log(f"错误: {str(e)}", "error")
            messagebox.showerror("错误", error_msg)
    
    def read_docx(self, filepath):
        """读取 .docx 文件"""
        if not DOCX_AVAILABLE:
            return "错误: python-docx 库未安装，请运行: pip install python-docx"
        
        try:
            doc = docx.Document(filepath)
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            if not paragraphs:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            paragraphs.append(row_text)
            
            if not paragraphs:
                return "文档为空"
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            return f"读取 .docx 失败: {str(e)}"
    
    def read_doc(self, filepath):
        """读取 .doc 文件"""
        try:
            with zipfile.ZipFile(filepath, 'r') as zf:
                if 'word/document.xml' in zf.namelist():
                    with zf.open('word/document.xml') as f:
                        content = f.read().decode('utf-8')
                        text = re.sub(r'<[^>]+>', ' ', content)
                        text = re.sub(r'\s+', ' ', text).strip()
                        if text:
                            return text
        except:
            pass
        
        return "无法读取 .doc 格式，请另存为 .docx 格式"
    
    def clear_content(self):
        """清空内容"""
        self.text_area.delete(1.0, tk.END)
        self.file_label.config(text="请选择一个Word文件...", fg="gray")
        self.status_label.config(text="就绪", fg="gray")
        self.current_file = None
        self.current_content = None
        self.log("内容已清空")
    
    def select_clean_file(self):
        """选择要被精简的Word文件"""
        file_types = [
            ("Word文档", "*.docx *.doc"),
            ("Word 2007+", "*.docx"),
            ("Word 97-2003", "*.doc"),
            ("所有文件", "*.*")
        ]
        
        filename = filedialog.askopenfilename(
            title="选择要被精简的Word文档",
            filetypes=file_types,
            initialdir=os.path.expanduser("~/Desktop")
        )
        
        if filename:
            self.clean_file = filename
            # 显示文件名（截取）
            basename = os.path.basename(filename)
            if len(basename) > 15:
                basename = basename[:12] + "..."
            self.clean_file_label.config(text=basename, fg="black")
            self.log(f"已选择精简文件: {filename}")
    
    def run_clean(self):
        """执行Word文档精简"""
        if self.clean_running:
            self.log("精简正在进行中，请稍候...", "warning")
            return
        
        if not self.clean_file:
            self.log("请先选择一个Word文件", "warning")
            messagebox.showwarning("提示", "请先选择要被精简的Word文件")
            return
        
        if not self.clean_file.endswith('.docx'):
            self.log("仅支持 .docx 格式的精简", "warning")
            messagebox.showwarning("提示", "仅支持 .docx 格式的精简")
            return
        
        result = messagebox.askyesno("确认", "确定要精简该文档吗？\n\n将合并同一人连续说的多段话，原文件保留。")
        if not result:
            return
        
        self.clean_running = True
        self.clean_btn.config(state=tk.DISABLED, text="精简中...")
        
        # 在线程中执行精简
        thread = threading.Thread(target=self._clean_worker)
        thread.daemon = True
        thread.start()
    
    def _clean_worker(self):
        """精简工作线程"""
        try:
            self.log("=" * 60, "title")
            self.log("开始执行文档精简...", "info")
            
            # 读取原文件
            doc = docx.Document(self.clean_file)
            
            # 精简处理：合并同一人连续说的多段话
            merged_count = 0
            
            # 获取所有非空段落
            paragraphs = [para for para in doc.paragraphs if para.text.strip()]
            
            if not paragraphs:
                self.log("文档为空", "warning")
                return
            
            # 新段落列表
            new_paragraphs = []
            
            i = 0
            while i < len(paragraphs):
                current_text = paragraphs[i].text.strip()
                
                # 检查是否是说话人格式：名字(时间戳): 内容
                # 匹配格式：XXX(HH:MM:SS): 内容
                import re
                match = re.match(r'^(.+?\(\d{2}:\d{2}:\d{2}\)):(.+)$', current_text)
                
                if match:
                    # 这是一个说话段落
                    current_speaker_with_time = match.group(1)  # 如: 阿标(00:00:06)
                    current_content = match.group(2).strip()
                    
                    # 提取说话人名字（去掉时间戳部分）
                    # 格式: 阿标(00:00:06) -> 名字是括号前的部分
                    current_name = current_speaker_with_time.split('(')[0] if '(' in current_speaker_with_time else ""
                    
                    # 查找后续连续的同一人说的段落
                    j = i + 1
                    while j < len(paragraphs):
                        next_text = paragraphs[j].text.strip()
                        next_match = re.match(r'^(.+?\(\d{2}:\d{2}:\d{2}\)):(.+)$', next_text)
                        
                        if next_match:
                            next_speaker_with_time = next_match.group(1)
                            next_content = next_match.group(2).strip()
                            
                            # 提取下一个说话人名字
                            next_name = next_speaker_with_time.split('(')[0] if '(' in next_speaker_with_time else ""
                            
                            # 判断是否是同一人
                            if current_name and next_name and current_name == next_name:
                                # 同一人，继续合并
                                current_content += next_content
                                merged_count += 1
                                j += 1
                            else:
                                break
                        else:
                            break
                    
                    # 添加合并后的段落
                    new_paragraphs.append(f"{current_speaker_with_time}:{current_content}")
                    
                    i = j
                else:
                    # 非说话段落，保留原样
                    new_paragraphs.append(current_text)
                    i += 1
            
            # 清空文档并写入新段落
            for para in doc.paragraphs:
                para._element.getparent().remove(para._element)
                para._element.getprevious()
            
            # 添加新段落
            # 先清空文档原有内容
            for para in doc.paragraphs[:]:
                p = para._element
                p.getparent().remove(p)
            
            # 添加新段落
            for text in new_paragraphs:
                doc.add_paragraph(text)
            
            # 生成新文件名
            dir_path = os.path.dirname(self.clean_file)
            basename = os.path.basename(self.clean_file)
            name_without_ext = os.path.splitext(basename)[0]
            ext = os.path.splitext(basename)[1]
            new_filename = f"{name_without_ext}_精简{ext}"
            new_filepath = os.path.join(dir_path, new_filename)
            
            # 保存新文件
            self.log(f"正在保存精简文件: {new_filename}...", "info")
            doc.save(new_filepath)
            
            # 输出结果
            self.log("=" * 60, "title")
            self.log("🎉 文档精简完成！", "success")
            self.log("-" * 60, "title")
            self.log(f"📊 合并段落: {merged_count} 处", "highlight")
            self.log(f"📄 原始段落: {len(paragraphs)} 段", "highlight")
            self.log(f"📄 精简后: {len(new_paragraphs)} 段", "highlight")
            self.log("-" * 60, "title")
            self.log(f"📁 新文件路径:", "info")
            self.log(f"   {new_filepath}", "highlight")
            self.log("=" * 60, "title")
            
            self.status_label.config(text=f"精简完成 - 合并 {merged_count} 处", fg="green")
            
            # 弹出提示
            self.root.after(0, lambda: messagebox.showinfo("✅ 精简完成", 
                f"🎉 文档精简完成！\n\n"
                f"📊 合并段落: {merged_count} 处\n"
                f"📄 原始: {len(paragraphs)} 段 → 精简后: {len(new_paragraphs)} 段\n\n"
                f"📁 新文件已保存至:\n{new_filepath}"))
            
        except Exception as e:
            self.log(f"精简失败: {str(e)}", "error")
            import traceback
            self.log(traceback.format_exc(), "error")
            messagebox.showerror("错误", f"精简失败: {str(e)}")
        
        finally:
            self.clean_running = False
            self.root.after(0, lambda: self.clean_btn.config(state=tk.NORMAL, text="✂️ 执行精简"))
    
    def load_library(self):
        """加载知识库"""
        library = {}
        
        if not os.path.exists(self.library_path):
            self.log(f"知识库文件不存在: {self.library_path}", "error")
            return library
        
        try:
            with open(self.library_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 解析 错误词=正确词 格式
            pattern = r'([^=\s]+)=([^=\s]+)'
            matches = re.findall(pattern, content)
            
            for wrong, right in matches:
                library[wrong] = right
            
            self.log(f"知识库已加载: {len(library)} 条词条")
            
        except Exception as e:
            self.log(f"加载知识库失败: {str(e)}", "error")
        
        return library
    
    def run_calibrate(self):
        """执行校准"""
        if self.calibration_running:
            self.log("校准正在进行中，请稍候...", "warning")
            return
        
        if not self.current_file:
            self.log("请先选择一个Word文件", "warning")
            messagebox.showwarning("提示", "请先选择一个Word文件")
            return
        
        if not self.current_file.endswith('.docx'):
            self.log("仅支持 .docx 格式的校准", "warning")
            messagebox.showwarning("提示", "仅支持 .docx 格式的校准")
            return
        
        result = messagebox.askyesno("确认", "确定要对文档进行错别字校准吗？\n\n将在原文件同位置生成新文件，原文件保留。")
        if not result:
            return
        
        self.calibration_running = True
        self.calibrate_btn.config(state=tk.DISABLED, text="校准中...")
        
        # 在线程中执行校准
        thread = threading.Thread(target=self._calibrate_worker)
        thread.daemon = True
        thread.start()
    
    def _calibrate_worker(self):
        """校准工作线程"""
        try:
            # 记录开始时间
            start_time = time.time()
            self.log("=" * 60, "title")
            self.log("开始执行校准...", "info")
            
            # 加载知识库
            library = self.load_library()
            if not library:
                self.log("知识库为空，无法执行校准", "error")
                return
            
            self.log(f"知识库词条数: {len(library)}")
            
            # 读取原文件
            doc = docx.Document(self.current_file)
            
            # 统计
            total_replacements = 0
            replacements_detail = []
            line_num = 0
            
            # 处理段落
            self.log("开始处理段落...", "info")
            for para in doc.paragraphs:
                line_num += 1
                original_text = para.text
                modified_text = original_text
                
                for wrong, right in library.items():
                    if wrong in modified_text:
                        count = modified_text.count(wrong)
                        modified_text = modified_text.replace(wrong, right)
                        total_replacements += count
                        for _ in range(count):
                            replacements_detail.append((line_num, wrong, right))
                
                if modified_text != original_text:
                    para.text = modified_text
            
            self.log(f"段落处理完成，共 {line_num} 段")
            
            # 处理表格
            self.log("开始处理表格...", "info")
            table_count = 0
            table_line_start = line_num
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    line_num += 1
                    for cell in row.cells:
                        original_text = cell.text
                        modified_text = original_text
                        
                        for wrong, right in library.items():
                            if wrong in modified_text:
                                count = modified_text.count(wrong)
                                modified_text = modified_text.replace(wrong, right)
                                total_replacements += count
                                for _ in range(count):
                                    replacements_detail.append((line_num, wrong, right))
                        
                        if modified_text != original_text:
                            cell.text = modified_text
            
            self.log(f"表格处理完成，共 {table_count} 个表格")
            
            # 计算时间
            elapsed_time = time.time() - start_time
            
            # 生成新文件名
            dir_path = os.path.dirname(self.current_file)
            basename = os.path.basename(self.current_file)
            name_without_ext = os.path.splitext(basename)[0]
            ext = os.path.splitext(basename)[1]
            new_filename = f"{name_without_ext}_校准{ext}"
            new_filepath = os.path.join(dir_path, new_filename)
            
            # 保存新文件
            self.log(f"正在保存文件: {new_filename}...", "info")
            doc.save(new_filepath)
            
            # 输出结果
            self.log("=" * 60, "title")
            self.log("🎉 校准处理完毕！", "success")
            self.log("-" * 60, "title")
            self.log(f"📊 处理词数: {total_replacements} 个", "highlight")
            self.log(f"⏱️ 总耗时: {elapsed_time:.2f} 秒", "highlight")
            self.log("-" * 60, "title")
            self.log(f"📁 新文件路径:", "info")
            self.log(f"   {new_filepath}", "highlight")
            self.log("=" * 60, "title")
            
            # 显示替换详情（全部显示）
            if replacements_detail:
                self.log("替换详情:", "info")
                for ln, wrong, right in replacements_detail:
                    self.log_replace(ln, wrong, right)
            
            # 最后显示总结信息
            self.log("=" * 60, "title")
            self.log("📊 处理词数: " + str(total_replacements) + " 个", "highlight")
            self.log("⏱️ 总耗时: " + str(elapsed_time) + " 秒", "highlight")
            self.log("📁 新文件: " + new_filepath, "highlight")
            self.log("=" * 60, "title")
            
            self.status_label.config(text=f"校准完成 - {total_replacements} 处替换", fg="green")
            
            # 弹出提示
            self.root.after(0, lambda: messagebox.showinfo("✅ 校准完成", 
                f"🎉 校准处理完毕！\n\n"
                f"📊 处理词数: {total_replacements} 个\n"
                f"⏱️ 总耗时: {elapsed_time:.2f} 秒\n\n"
                f"📁 新文件已保存至:\n{new_filepath}"))
            
        except Exception as e:
            self.log(f"校准失败: {str(e)}", "error")
            messagebox.showerror("错误", f"校准失败: {str(e)}")
        
        finally:
            self.calibration_running = False
            self.root.after(0, lambda: self.calibrate_btn.config(state=tk.NORMAL, text="🔧 执行校准"))
    
    def open_library(self):
        """打开知识库"""
        self.log("点击了【知识库】按钮")
        base_dir = os.path.dirname(os.path.abspath(__file__))
        lib_path = os.path.join(base_dir, "errorLibrary", "readLibrary.py")
        if os.path.exists(lib_path):
            try:
                import subprocess
                subprocess.Popen([sys.executable, lib_path])
                self.log("已打开知识库")
            except Exception as e:
                self.log(f"打开知识库失败: {e}")
        else:
            self.log(f"知识库文件不存在: {lib_path}")
    
    def open_settings(self):
        """打开设置"""
        self.log("点击了【设置】按钮")
        messagebox.showinfo("⚙️ 功能开发中", "AI让我们觉醒加速！")


if __name__ == "__main__":
    # 检查环境
    if not DOCX_AVAILABLE:
        print("警告: python-docx 未安装")
        print("请运行: pip install python-docx")
        input("按回车键退出...")
    
    try:
        app = WordReaderApp()
    except Exception as e:
        print(f"启动失败: {e}")
        input("按回车键退出...")
